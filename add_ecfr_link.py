import re
from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


def add_hyperlink(paragraph, url, text):
    """
    Adds a hyperlink to a paragraph in a docx file.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Style the link (Standard Blue + Underline)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def process_document(input_file, output_file):
    doc = Document(input_file)

    # REGEX EXPLANATION:
    # 97\.             -> Matches literal "97."
    # (\d+)            -> Group 2: Captures the section number (e.g., "111") for the base URL
    # (?:\([a-zA-Z0-9]+\))* -> Non-capturing group that repeats. Matches subsections like (a), (1), (ii).
    #
    # This matching pattern (Group 0) is used for the anchor tag #p-97.111(a)(1)
    pattern = re.compile(r"(97\.(\d+)(?:\([a-zA-Z0-9]+\))*)")

    for paragraph in doc.paragraphs:
        # We only process paragraphs containing "97." to save time
        if "97." in paragraph.text:
            original_text = paragraph.text

            # Find all citations in the text
            matches = list(pattern.finditer(original_text))

            # If no regex matches found (false positive on "97."), skip
            if not matches:
                continue

            # Clear the paragraph to rebuild it with links
            paragraph.clear()

            last_pos = 0
            for match in matches:
                start, end = match.span()

                # 1. Append text occurring BEFORE the link (e.g., "[", ", ", or normal text)
                if start > last_pos:
                    paragraph.add_run(original_text[last_pos:start])

                # 2. Extract Data
                full_citation = match.group(1)  # e.g., "97.111(a)(1)"
                section_num = match.group(2)  # e.g., "111"

                # 3. Construct the deep-link URL
                # Format: .../part-97/section-97.111#p-97.111(a)(1)
                url = (
                    f"https://www.ecfr.gov/current/title-47/part-97/"
                    f"section-97.{section_num}#p-{full_citation}"
                )

                # 4. Add the Hyperlink
                add_hyperlink(paragraph, url, full_citation)

                last_pos = end

            # 5. Append any remaining text after the last link (e.g., "]")
            if last_pos < len(original_text):
                paragraph.add_run(original_text[last_pos:])

    doc.save(output_file)
    print(f"Success! Saved to: {output_file}")


# Usage
process_document(
    "2026-2030 Technician Pool and Syllabus Public Release Dec 18 2025.docx",
    "2026_Technician_Pool_Linked.docx",
)
